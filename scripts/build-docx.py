# -*- coding: utf-8 -*-
"""Build SDLC Lean guide .docx from repo content. Requires: python-docx."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "docs", "SDLC-Lean-Guide.docx")

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x59, 0x59, 0x59)

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = ACCENT
    return p


def para(text="", bold=False, italic=False, size=None, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    return p


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t


# ---------------------------------------------------------------- Cover
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("SDLC Lean Workflow")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Hướng dẫn quy trình phát triển tinh gọn cho team AI-assisted")
r.italic = True
r.font.size = Pt(13)
r.font.color.rgb = GREY

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Version 1.3.0  •  Graphify Enabled  •  MIT License").font.size = Pt(10)

doc.add_paragraph()
para("Nguyên tắc cốt lõi:", bold=True)
para("Mọi thay đổi phải nhỏ, truy vết được, được review, được test, và được báo cáo "
     "(small, traceable, reviewed, tested, reported).", italic=True)

doc.add_page_break()

# ---------------------------------------------------------------- 1. Tổng quan
h("1. Tổng quan", 1)
para("SDLC Lean là một framework quy trình phát triển nhẹ (low-ceremony) dùng cho team "
     "phát triển phần mềm có hỗ trợ AI. Dùng repo này như một template SDLC tái sử dụng cho:")
for x in ["Phát triển tính năng mới (new feature)",
          "Sửa lỗi (bug fixing)",
          "Chỉnh sửa nhỏ API / business logic",
          "Refactor / modification",
          "Phát triển với AI-agent + Graphify impact search"]:
    bullet(x)

# ---------------------------------------------------------------- 2. Yêu cầu & Cài đặt
h("2. Yêu cầu & Cài đặt", 1)
para("Yêu cầu môi trường:", bold=True)
bullet("Node.js >= 18 (chạy CLI sdlc-workflow).")
bullet("Git (quản lý source + PR).")
bullet("PowerShell (chạy script validation trên Windows).")
bullet("Tùy chọn: Python 3 + python-docx (chỉ cần khi build file .docx này).")

para()
para("Cách 1 — Scaffold vào project sẵn có:", bold=True)
para("Chạy tại thư mục gốc của project bạn muốn áp dụng workflow:")
code("npx sdlc-workflow init")
para("Lệnh này tạo: docs/, templates/, adr/, .github/, scripts/ và rule cho agent "
     "(.cursor/rules/, .claude/CLAUDE.md, AGENTS.md, .agents/skills/).", italic=True, color=GREY)

para()
para("Cách 2 — Cài global skills cho agent (Cursor, Codex):", bold=True)
code("npx sdlc-workflow install")
para("Cài vào home dir: ~/.cursor/skills/, ~/.codex/AGENTS.md, ~/.agents/skills/.",
     italic=True, color=GREY)

para()
para("Cách 3 — Clone & chạy trực tiếp từ repo:", bold=True)
code("git clone <repo-url> sdlc-lean\n"
     "cd sdlc-lean\n"
     "node bin/cli.js version")

para()
para("Kiểm tra cài đặt:", bold=True)
code("npx sdlc-workflow version\n"
     "powershell -ExecutionPolicy Bypass -File scripts/validate-workflow.ps1")

# ---------------------------------------------------------------- 3. Hướng dẫn sử dụng
h("3. Hướng dẫn sử dụng", 1)
para("Luồng dùng hằng ngày cho một task code:", bold=True)

para()
para("Bước 1 — Mở một task (epic) mới:", bold=True)
para("Trong AI agent (Claude Code / Cursor) dùng slash command:")
code('/sdlc-lean Tôi muốn <goal> | Dữ liệu đầu vào <input> | Kết quả mong muốn <output>')
para("Hoặc chạy CLI trực tiếp:")
code('node bin/cli.js new "Tôi muốn tạo API login | Input: email+password | Output: JWT token"')
para("→ Tạo work/<NNN>-<slug>/ kèm bộ file evidence, và set epic đó thành active "
     "(ghi vào work/.active).", italic=True, color=GREY)

para()
para("Bước 2 — Làm việc trong epic (vòng 6 bước):", bold=True)
para("Agent điền epic-brief.md rồi chạy lần lượt 6 bước (xem mục 4). Trong khi epic còn "
     "active, cứ tiếp tục làm trong đó — KHÔNG cần chạy lại /sdlc-lean. Mọi evidence "
     "(impact plan, test, build, review) ghi vào folder epic.")

para()
para("Bước 3 — Đóng task khi hoàn thành:", bold=True)
code("/finish")
para("hoặc:")
code("node bin/cli.js finish")
para("→ Đánh dấu epic done, sinh SUMMARY.md, xóa con trỏ active. Request mới sẽ mở epic mới.",
     italic=True, color=GREY)

para()
para("Mẹo dùng tốt:", bold=True)
bullet("Luôn viết request theo mẫu Tôi muốn | Input | Output để goal/input/output rõ ràng.")
bullet("Một task = một epic = một folder; đừng trộn nhiều mục tiêu vào một epic.")
bullet("Dừng ở Approval Gate — không để agent sửa code lớn trước khi bạn duyệt.")
bullet("Trước khi mở PR, đối chiếu Definition of Done (mục 10).")

# ---------------------------------------------------------------- 4. Quy trình 6 bước
h("4. Quy trình 6 bước", 1)
para("Khi được yêu cầu sửa hoặc thêm code, thực hiện chặt chẽ 6 bước:")

steps = [
    ("Clarify (Làm rõ) — ",
     "Hỏi 1-3 câu hỏi trọng tâm và phân loại công việc (Feature / Refactor / Bugfix / "
     "Small change / API change / Business logic). Với Bugfix, hỏi bước tái hiện hoặc error log."),
    ("Search with Graphify (Tìm kiếm) — ",
     "Dùng Graphify để định vị file, function, class, API, dependency và test liên quan. "
     "Nếu Graphify không có, nêu rõ và dùng repo/IDE search làm fallback."),
    ("Impact Plan (Kế hoạch tác động) — ",
     "Liệt kê: file cần đổi, thay đổi dự kiến, side effects, test/build cần chạy, mức rủi ro, đường rollback."),
    ("Approval Gate (Cổng phê duyệt) — ",
     "DỪNG trước các sửa đổi không tầm thường và chờ phê duyệt rõ ràng. Đây là điểm dừng cứng."),
    ("Code + Test + Build — ",
     "Triển khai theo plan đã duyệt, chạy test/build/validation. Chỉ retry lỗi do thay đổi hiện tại, "
     "tối đa 3 lần rồi báo blocker."),
    ("Report (Báo cáo) — ",
     "Tóm tắt: file đã sửa, trạng thái test, trạng thái build/validation, check bị skip, rủi ro còn lại."),
]
for prefix, body in steps:
    numbered(body, bold_prefix=prefix)

para()
para("Approval gate là điểm dừng cứng: không sửa code không tầm thường trước khi được duyệt. "
     "Mức rủi ro (risk-matrix) quyết định độ mạnh của phê duyệt.", italic=True, color=GREY)

# ---------------------------------------------------------------- 3. Output mỗi bước
h("5. Output bắt buộc theo từng bước", 1)
table(["Bước", "Output"],
      [["Clarify", "change type, scope, acceptance criteria"],
       ["Search", "file / function / test liên quan"],
       ["Impact", "files, side effects, risk, rollback"],
       ["Approval", "phê duyệt rõ ràng trước khi sửa"],
       ["Execution", "code + tests/build"],
       ["Report", "evidence, skipped checks, risks"]])

# ---------------------------------------------------------------- 4. Change types
h("6. Các loại thay đổi (Change Types)", 1)
table(["Loại", "Dùng khi", "Tài liệu bắt buộc"],
      [["Feature", "Năng lực mới", "requirement brief, impact plan, tests"],
       ["Bugfix", "Sửa defect", "repro, root cause, regression test"],
       ["Small change", "Sửa nhỏ an toàn", "short impact note, smoke test"],
       ["API change", "Đổi endpoint/contract", "compatibility, contract note, integration test"],
       ["Business logic", "Đổi rule/quyết định", "rule source, edge cases, business tests"],
       ["Refactor", "Đổi cấu trúc nội bộ", "bằng chứng no behavior-change"]])

# ---------------------------------------------------------------- 5. Epic workflow
h("7. Epic Workflow (một thư mục mỗi task)", 1)
para("Mỗi task là một \"epic\" lưu tại work/<NNN>-<slug>/ với brief + file evidence riêng. "
     "Epic đang hoạt động được theo dõi trong work/.active.")
bullet("Tạo task mới — sets active, rồi chạy vòng 6 bước, ghi evidence vào folder epic.",
       bold_prefix="/sdlc-lean <request>  (hoặc  node bin/cli.js new \"<request>\")  ")
bullet("Trong khi epic active, tiếp tục làm trong đó — không cần chạy lại /sdlc-lean.",
       bold_prefix="Tiếp tục:  ")
bullet("Đánh dấu done, ghi SUMMARY.md, xóa con trỏ active.",
       bold_prefix="/finish  (hoặc  node bin/cli.js finish)  ")

para()
para("Quy tắc prompt:", bold=True)
code('Tôi muốn <goal> | Dữ liệu đầu vào <input> | Kết quả mong muốn <output>')
para("để goal, input và output đều rõ ràng.", italic=True, color=GREY)

# ---------------------------------------------------------------- 6. CLI
h("8. Lệnh CLI", 1)
table(["Lệnh", "Tác dụng"],
      [["npx sdlc-workflow init", "Scaffold docs, templates, agent rules vào project"],
       ["npx sdlc-workflow install", "Cài global skills (Cursor, Codex) vào home dir"],
       ['npx sdlc-workflow new "<text>"', "Bắt đầu task: tạo work/<NNN>-<slug>/ + set active"],
       ["npx sdlc-workflow finish", "Đóng task active: mark done + ghi SUMMARY.md"],
       ["npx sdlc-workflow version", "In version hiện tại"]])

# ---------------------------------------------------------------- 7. Quickstart
h("9. Quickstart cho team", 1)
for x in ["Đọc docs/workflow.md.",
          "Chọn change-flow doc phù hợp.",
          "Điền templates/requirement-brief.md hoặc issue template.",
          "Chạy Graphify/code search và điền templates/impact-analysis.md.",
          "Lấy phê duyệt (approval).",
          "Triển khai và đính kèm test/build evidence.",
          "Mở PR dùng .github/PULL_REQUEST_TEMPLATE.md."]:
    numbered(x)

# ---------------------------------------------------------------- 8. Definition of Done
h("10. Definition of Done tối thiểu", 1)
for x in ["requirement rõ ràng",
          "đã chọn change type",
          "Graphify/code search xong hoặc skip có nêu rõ",
          "impact plan đã được duyệt",
          "test/build đã chạy hoặc skip kèm lý do",
          "PR có evidence và rollback note",
          "có reviewer approval theo mức rủi ro"]:
    bullet(x)

# ---------------------------------------------------------------- 9. Repo map
h("11. Bản đồ repository", 1)
code("docs/                          SDLC policies, flow docs, checklists\n"
     "templates/                     Reusable planning/evidence templates\n"
     "adr/                           Architecture decision records\n"
     ".github/                       PR/issue templates + validation workflow\n"
     "scripts/validate-workflow.ps1  Lightweight structure/PR validation\n"
     "bin/cli.js                     SDLC workflow CLI (init/new/finish/install)\n"
     ".claude/CLAUDE.md              Claude Code project instructions")

# ---------------------------------------------------------------- 10. Adoption + rules
h("12. Mức độ áp dụng & Quy tắc an toàn", 1)
table(["Level", "Ý nghĩa"],
      [["Basic", "Dùng docs + PR template thủ công"],
       ["Team", "Bắt buộc issue/PR templates + validation script"],
       ["Professional", "Thêm branch protection, CODEOWNERS, CI/test/build gates thực"]])

para()
para("Quy tắc an toàn (non-negotiables):", bold=True)
for x in ["Giữ thay đổi tinh gọn và đúng scope.",
          "Không che giấu check Graphify/test/build bị skip.",
          "Không tiếp tục retry loop cho lỗi không liên quan thay đổi hiện tại.",
          "Không merge, push, publish, delete, hay hành động không thể đảo ngược trừ khi được yêu cầu rõ ràng.",
          "Không sửa code không tầm thường trước khi được duyệt.",
          "Không merge khi chưa có review evidence."]:
    bullet(x)

# ------------------------------------------------------ 13. Database safety
h("13. IMPORTANT — An toàn Database", 1)
para("Đây là quy tắc bắt buộc (non-negotiable), áp dụng cho mọi agent và mọi loại thay đổi:",
     bold=True, color=RGBColor(0xC0, 0x00, 0x00))
bullet("KHÔNG dùng raw SQL để thao tác database. Luôn đi qua ORM / query builder / "
       "repository layer của project.", bold_prefix="• ")
bullet("KHÔNG dùng test, script, migration, seed hay bất kỳ cơ chế nào để "
       "DELETE / DROP / UPDATE / TRUNCATE dữ liệu database của user.", bold_prefix="• ")
bullet("Mọi thao tác chạm tới dữ liệu thật là irreversible action → cấm trừ khi user "
       "yêu cầu rõ ràng từng lần.", bold_prefix="• ")
para()
para("Quy tắc này được nhúng trong .claude/CLAUDE.md, AGENTS.md, .cursor/rules/ và các "
     "skill agent — nên agent sẽ tự tuân thủ trong suốt vòng 6 bước.", italic=True, color=GREY)

# ---------------------------------------------------------------- 14. Automation
h("14. Tự động hóa (Validation)", 1)
para("Chạy cục bộ:")
code("powershell -ExecutionPolicy Bypass -File scripts/validate-workflow.ps1")
para("GitHub Actions chạy cùng validation trên push/PR qua "
     ".github/workflows/workflow-validation.yml.")

# ---------------------------------------------------------------- Footer
doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("SDLC Lean Workflow • Generated from repo docs • MIT License")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = GREY

doc.save(OUT)
print("Saved:", OUT)
